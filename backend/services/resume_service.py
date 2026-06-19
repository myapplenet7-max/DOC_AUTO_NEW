"""
Resume DOCX generation service.

Supports three resume types:
  - fresher:     Simple, clean, education-first layout
  - experienced: Multi-section with work history and achievements
  - creative:    Bold header with accent colour, two-column skills
"""
import json
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("docauto.resume")

# ── Accent colours per resume type ───────────────────────────────────────────
ACCENT = {
    "fresher":     RGBColor(0x1a, 0x56, 0xdb),   # indigo-600
    "experienced": RGBColor(0x04, 0x7a, 0x5e),   # emerald-700
    "creative":    RGBColor(0x7c, 0x3a, 0xed),   # violet-600
}

def _hex_to_rgb(hex_str: str):
    """Convert #rrggbb to RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (OOXML shading)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _heading(doc: Document, text: str, accent: RGBColor, size_pt: int = 11):
    """Section heading with coloured bottom border."""
    p    = doc.add_paragraph()
    run  = p.add_run(text.upper())
    run.bold = True
    run.font.size  = Pt(size_pt)
    run.font.color.rgb = accent
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), str(accent))
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _bullet(doc: Document, text: str, indent: float = 0.25):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_after   = Pt(1)
    p.paragraph_format.space_before  = Pt(0)
    return p


def _watermark_paragraph(doc: Document):
    p   = doc.add_paragraph()
    run = p.add_run("━━━━━━  PREVIEW ONLY — NOT FOR DISTRIBUTION — DOCAUTO  ━━━━━━")
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xd1, 0x1a, 0x2c)
    run.font.bold  = True
    p.alignment    = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _name_header(doc: Document, personal: dict, accent: RGBColor):
    """Big name + contact line at the top."""
    name_p = doc.add_paragraph()
    name_r = name_p.add_run(personal.get("name", "Your Name"))
    name_r.bold = True
    name_r.font.size  = Pt(22)
    name_r.font.color.rgb = accent
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)

    # contact line
    parts = []
    for key in ("phone", "email", "location", "linkedin", "portfolio"):
        v = personal.get(key, "").strip()
        if v:
            parts.append(v)
    contact_p = doc.add_paragraph(" · ".join(parts))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    for run in contact_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)


def _education_section(doc: Document, entries: list, accent: RGBColor):
    if not entries:
        return
    _heading(doc, "Education", accent)
    for e in entries:
        institution = e.get("institution", "")
        degree = e.get("degree", "")
        field  = e.get("field", "")
        year   = e.get("year", "")
        grade  = e.get("grade", "")

        p    = doc.add_paragraph()
        r    = p.add_run(institution)
        r.bold = True
        r.font.size = Pt(10)
        if year:
            yr = p.add_run(f"  {year}")
            yr.font.size = Pt(9)
            yr.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        sub_parts = []
        if degree: sub_parts.append(degree)
        if field:  sub_parts.append(field)
        if grade:  sub_parts.append(f"Grade: {grade}")
        if sub_parts:
            sp = doc.add_paragraph(", ".join(sub_parts))
            for run in sp.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            sp.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_after = Pt(0)


def _experience_section(doc: Document, entries: list, accent: RGBColor):
    if not entries:
        return
    _heading(doc, "Work Experience", accent)
    for e in entries:
        company = e.get("company", "")
        title   = e.get("title", "")
        start   = e.get("start", "")
        end     = e.get("end", "Present") if not e.get("current") else "Present"
        desc    = e.get("description", "")

        p   = doc.add_paragraph()
        r   = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        if company:
            c = p.add_run(f"  ·  {company}")
            c.font.size = Pt(10)
            c.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        if start:
            period = f"  {start} – {end}"
            t = p.add_run(period)
            t.font.size = Pt(9)
            t.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        if desc:
            for line in desc.split("\n"):
                line = line.strip()
                if line:
                    _bullet(doc, line)
        p.paragraph_format.space_after = Pt(2)


def _skills_section(doc: Document, skills: list, accent: RGBColor):
    if not skills:
        return
    _heading(doc, "Skills", accent)
    skill_text = " · ".join(s.strip() for s in skills if s.strip())
    p = doc.add_paragraph(skill_text)
    for run in p.runs:
        run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _certifications_section(doc: Document, certs: list, accent: RGBColor):
    if not certs:
        return
    _heading(doc, "Certifications", accent)
    for c in certs:
        name   = c.get("name", "")
        issuer = c.get("issuer", "")
        year   = c.get("year", "")
        parts  = [name]
        if issuer: parts.append(issuer)
        if year:   parts.append(year)
        _bullet(doc, " · ".join(p for p in parts if p))


def _projects_section(doc: Document, projects: list, accent: RGBColor):
    if not projects:
        return
    _heading(doc, "Projects", accent)
    for proj in projects:
        name = proj.get("name", "")
        desc = proj.get("description", "")
        tech = proj.get("tech", "")
        url  = proj.get("url", "")

        p   = doc.add_paragraph()
        r   = p.add_run(name)
        r.bold = True
        r.font.size = Pt(10)
        if tech:
            t = p.add_run(f"  [{tech}]")
            t.font.size = Pt(9)
            t.font.color.rgb = accent
        if desc:
            d = doc.add_paragraph(desc)
            for run in d.runs:
                run.font.size = Pt(9)
            d.paragraph_format.left_indent = Inches(0.2)
            d.paragraph_format.space_after = Pt(4)
        if url:
            u = doc.add_paragraph(url)
            for run in u.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = accent
            u.paragraph_format.left_indent = Inches(0.2)
            u.paragraph_format.space_after = Pt(4)


def _achievements_section(doc: Document, achievements: list, accent: RGBColor):
    if not achievements:
        return
    _heading(doc, "Achievements", accent)
    for a in achievements:
        a = a.strip()
        if a:
            _bullet(doc, a)


def generate_resume_docx(data: dict, output_path: str, watermark: bool = False):
    """
    Generate a resume DOCX from `data` dict.
    If watermark=True, inserts visible preview watermark lines.
    """
    resume_type = data.get("type", "fresher")
    personal    = data.get("personal", {})
    summary     = data.get("summary", "")
    education   = data.get("education", [])
    experience  = data.get("experience", [])
    skills      = data.get("skills", [])
    certs       = data.get("certifications", [])
    projects    = data.get("projects", [])
    achievements = data.get("achievements", [])
    accent_hex  = data.get("accent_color", "")

    accent = _hex_to_rgb(accent_hex) if accent_hex else ACCENT.get(resume_type, ACCENT["fresher"])

    doc = Document()

    # ── page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── default font ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Watermark banner (top) ────────────────────────────────────────────────
    if watermark:
        _watermark_paragraph(doc)

    # ── Name & Contact ────────────────────────────────────────────────────────
    _name_header(doc, personal, accent)

    # ── Type-specific layout ──────────────────────────────────────────────────
    if resume_type == "fresher":
        # Objective/Summary → Education → Projects → Skills → Certifications
        if summary:
            _heading(doc, "Objective / Summary", accent)
            p = doc.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        _education_section(doc, education, accent)
        _projects_section(doc, projects, accent)
        _skills_section(doc, skills, accent)
        _certifications_section(doc, certs, accent)

    elif resume_type == "experienced":
        # Summary → Experience → Education → Skills → Certifications → Achievements
        if summary:
            _heading(doc, "Professional Summary", accent)
            p = doc.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        _experience_section(doc, experience, accent)
        _education_section(doc, education, accent)
        _skills_section(doc, skills, accent)
        _certifications_section(doc, certs, accent)
        _achievements_section(doc, achievements, accent)

    elif resume_type == "creative":
        # Summary → Skills (bold tags) → Experience → Projects → Education → Certifications
        if summary:
            _heading(doc, "About Me", accent)
            p = doc.add_paragraph(summary)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True
            p.paragraph_format.space_after = Pt(4)

        # Skills as tags in a shaded paragraph
        if skills:
            _heading(doc, "Skills & Tools", accent)
            tag_para = doc.add_paragraph()
            for i, skill in enumerate(skills):
                s = skill.strip()
                if not s:
                    continue
                r = tag_para.add_run(f" {s} ")
                r.font.size = Pt(10)
                r.font.color.rgb = accent
                r.bold = True
                if i < len(skills) - 1:
                    sep = tag_para.add_run("  ")
                    sep.font.size = Pt(10)
            tag_para.paragraph_format.space_after = Pt(6)

        portfolio = personal.get("portfolio", "").strip()
        if portfolio:
            pp = doc.add_paragraph()
            r = pp.add_run("Portfolio / Work Samples: ")
            r.bold = True
            r.font.size = Pt(10)
            l = pp.add_run(portfolio)
            l.font.size = Pt(10)
            l.font.color.rgb = accent
            pp.paragraph_format.space_after = Pt(4)

        _experience_section(doc, experience, accent)
        _projects_section(doc, projects, accent)
        _education_section(doc, education, accent)
        _certifications_section(doc, certs, accent)

    # ── Watermark banner (bottom) ─────────────────────────────────────────────
    if watermark:
        _watermark_paragraph(doc)

    doc.save(output_path)
    logger.info("Resume saved: type=%s path=%s watermark=%s", resume_type, output_path, watermark)
    return output_path
