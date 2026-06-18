import re
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def extract_placeholder_keys(template_content: str) -> list:
    """Return all unique {{KEY}} placeholders found in the template."""
    return list(dict.fromkeys(re.findall(r'\{\{([A-Z0-9_]+)\}\}', template_content)))


def fill_template(template_content: str, fields: dict) -> str:
    """Replace {{KEY}} placeholders with actual values."""
    result = template_content
    for key, value in fields.items():
        result = result.replace(f"{{{{{key}}}}}", str(value) if value else f"[{key}]")
    # Leave unfilled placeholders visually marked
    result = re.sub(r'\{\{([A-Z0-9_]+)\}\}', r'[\1]', result)
    return result


def generate_filled_docx(template_content: str, fields: dict, output_path: str, template_name: str = "Document"):
    """Generate a DOCX from a template with filled values."""
    filled_text = fill_template(template_content, fields)

    doc = Document()
    _set_telugu_capable_font(doc)

    # Title
    title = doc.add_heading(template_name.upper(), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_paragraph("")

    # Write the filled content preserving paragraphs
    for para_text in filled_text.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        # Bold text in [UNFILLED] placeholders
        parts = re.split(r'(\[[A-Z0-9_]+\])', para_text)
        for part in parts:
            if re.match(r'\[[A-Z0-9_]+\]', part):
                run = p.add_run(part)
                run.bold = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                p.add_run(part)

    doc.save(output_path)


def generate_template_docx_preview(template_content: str, output_path: str, template_name: str = "Template"):
    """Generate a DOCX showing the template with placeholders highlighted."""
    doc = Document()
    _set_telugu_capable_font(doc)

    title = doc.add_heading(f"TEMPLATE: {template_name.upper()}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Created on: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_paragraph("")

    # Field list
    keys = extract_placeholder_keys(template_content)
    if keys:
        doc.add_heading("Required Fields", level=2)
        for key in keys:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{{{{{key}}}}}").bold = True
            p.add_run(f"  →  {key.replace('_', ' ').title()}")
        doc.add_paragraph("")

    doc.add_heading("Template Content", level=2)
    for para_text in template_content.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        parts = re.split(r'(\{\{[A-Z0-9_]+\}\})', para_text)
        for part in parts:
            if re.match(r'\{\{[A-Z0-9_]+\}\}', part):
                run = p.add_run(part)
                run.bold = True
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            else:
                p.add_run(part)

    doc.save(output_path)


def build_field_schema(template_content: str) -> str:
    """Build JSON schema listing all placeholder fields in a template."""
    keys = extract_placeholder_keys(template_content)
    schema = {
        "fields": [
            {
                "key": key,
                "label": key.replace("_", " ").title(),
                "type": _guess_field_type(key),
                "required": True,
            }
            for key in keys
        ]
    }
    return json.dumps(schema, ensure_ascii=False, indent=2)


def _guess_field_type(key: str) -> str:
    key_lower = key.lower()
    if "date" in key_lower or "dob" in key_lower:
        return "date"
    if "mobile" in key_lower or "phone" in key_lower:
        return "tel"
    if "email" in key_lower:
        return "email"
    if "amount" in key_lower or "price" in key_lower or "fee" in key_lower:
        return "currency"
    if "number" in key_lower or "no" == key_lower[-2:]:
        return "text"
    if "address" in key_lower or "boundaries" in key_lower or "property" in key_lower:
        return "textarea"
    return "text"


def _set_telugu_capable_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Nirmala UI"
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Nirmala UI")
    rfonts.set(qn("w:hAnsi"), "Nirmala UI")
    rfonts.set(qn("w:cs"), "Nirmala UI")
