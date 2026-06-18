"""
Format-preserving document processing.

For DOCX inputs: clones the original file and performs in-place text replacement,
preserving ALL formatting (fonts, sizes, bold, italic, colors, tables, headers, footers,
page layout, borders, images, etc.).

For other formats: falls back to structured text extraction.
"""
import os
import re
import json
import copy
from docx import Document
from docx.oxml.ns import qn


# ── In-place DOCX replacement ────────────────────────────────────────────────

def create_template_from_docx(
    source_path: str,
    detected_fields: dict,
    template_output_path: str,
) -> dict:
    """
    Clone source DOCX and replace detected field values with {{PLACEHOLDER}} tokens in-place.
    Returns {placeholder_key: original_value} mapping.
    Preserves 100% of original formatting.
    """
    import shutil
    shutil.copy2(source_path, template_output_path)

    doc = Document(template_output_path)

    # Build sorted replacement map (longest first to avoid partial matches)
    replacements = {}
    for field_key, value in detected_fields.items():
        if field_key == "raw_text":
            continue
        if value and str(value).strip():
            placeholder = f"{{{{{field_key.upper()}}}}}"
            replacements[str(value).strip()] = placeholder

    sorted_replacements = sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)

    _replace_in_document(doc, sorted_replacements)
    doc.save(template_output_path)
    return {v.strip("{}"): k for k, v in replacements.items()}


def fill_template_docx(
    template_path: str,
    fill_values: dict,
    output_path: str,
) -> None:
    """
    Clone template DOCX and replace {{PLACEHOLDER}} tokens with actual values.
    Preserves 100% of original formatting.
    """
    import shutil
    shutil.copy2(template_path, output_path)

    doc = Document(output_path)

    # Build replacement list: {{KEY}} → value
    replacements = []
    for key, value in fill_values.items():
        placeholder = f"{{{{{key.upper()}}}}}"
        replacements.append((placeholder, value if value else f"[{key}]"))

    # Also handle lowercase variants
    for key, value in fill_values.items():
        placeholder = f"{{{{{key}}}}}"
        replacements.append((placeholder, value if value else f"[{key}]"))

    sorted_replacements = sorted(replacements, key=lambda x: len(x[0]), reverse=True)

    _replace_in_document(doc, sorted_replacements)
    doc.save(output_path)


def _replace_in_document(doc: Document, sorted_replacements: list) -> None:
    """Apply replacements throughout all document structures."""
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, sorted_replacements)

    # Tables (including nested)
    for table in doc.tables:
        _replace_in_table(table, sorted_replacements)

    # Headers and footers
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                _replace_in_paragraph(para, sorted_replacements)
            for table in section.header.tables:
                _replace_in_table(table, sorted_replacements)
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                _replace_in_paragraph(para, sorted_replacements)
            for table in section.footer.tables:
                _replace_in_table(table, sorted_replacements)
        except Exception:
            pass


def _replace_in_table(table, sorted_replacements: list) -> None:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _replace_in_paragraph(para, sorted_replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, sorted_replacements)


def _replace_in_paragraph(para, sorted_replacements: list) -> None:
    """
    Replace text in a paragraph preserving run-level formatting.

    Strategy:
    1. Try single-run replacement first (preserves all run formatting perfectly).
    2. If old_text spans multiple runs, rebuild the full text on the first run
       that contains the start of the match, and clear the other affected runs.
    """
    if not para.runs:
        return

    for old_text, new_text in sorted_replacements:
        # Fast path: contained within a single run
        replaced_in_run = False
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced_in_run = True

        if not replaced_in_run:
            # Check if it spans multiple runs
            full = "".join(r.text for r in para.runs)
            if old_text in full:
                _cross_run_replace(para, old_text, new_text)


def _cross_run_replace(para, old_text: str, new_text: str) -> None:
    """
    Handle replacement where old_text spans multiple runs.
    Merges affected runs into the first run of the span, preserving its format.
    """
    runs = para.runs
    if not runs:
        return

    # Build cumulative offsets
    texts = [r.text for r in runs]
    full = "".join(texts)

    if old_text not in full:
        return

    start_idx = full.index(old_text)
    end_idx = start_idx + len(old_text)

    # Find which run indices are in [start_idx, end_idx)
    pos = 0
    run_spans = []
    for i, t in enumerate(texts):
        run_spans.append((pos, pos + len(t)))
        pos += len(t)

    affected_runs = []
    for i, (rs, re_) in enumerate(run_spans):
        if rs < end_idx and re_ > start_idx:
            affected_runs.append(i)

    if not affected_runs:
        return

    first_idx = affected_runs[0]
    new_full = full[:start_idx] + new_text + full[end_idx:]

    # Rebuild: give all text to first run, zero out others
    # We only rebuild the affected portion to avoid disturbing unaffected runs
    prefix = full[:run_spans[first_idx][0]]
    suffix = full[run_spans[affected_runs[-1]][1]:]

    # Text that goes into first affected run
    runs[first_idx].text = new_full[len(prefix): len(new_full) - len(suffix)]

    # Clear all other affected runs
    for i in affected_runs[1:]:
        runs[i].text = ""


# ── DOCX → HTML (for browser preview) ───────────────────────────────────────

def docx_to_preview_html(doc_path: str) -> str:
    """
    Convert a DOCX to HTML using mammoth for in-browser preview.
    Returns an HTML string.
    """
    try:
        import mammoth
        with open(doc_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value
        # Wrap in a styled container
        return _wrap_preview_html(html)
    except ImportError:
        return _docx_to_basic_html(doc_path)
    except Exception as e:
        return f"<div class='text-red-500'>Preview error: {e}</div>"


def _wrap_preview_html(inner_html: str) -> str:
    return f"""
<style>
  .doc-preview {{
    font-family: 'Nirmala UI', 'Noto Sans Telugu', Arial, sans-serif;
    font-size: 13px;
    line-height: 1.7;
    color: #1a1a1a;
    padding: 24px 32px;
    max-width: 680px;
    margin: 0 auto;
    background: white;
  }}
  .doc-preview h1, .doc-preview h2, .doc-preview h3 {{
    font-weight: bold;
    margin: 16px 0 8px 0;
  }}
  .doc-preview h1 {{ font-size: 18px; text-align: center; }}
  .doc-preview h2 {{ font-size: 15px; }}
  .doc-preview p {{ margin: 6px 0; }}
  .doc-preview table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
  .doc-preview td, .doc-preview th {{
    border: 1px solid #ccc;
    padding: 6px 10px;
    vertical-align: top;
  }}
  .doc-preview strong {{ font-weight: bold; }}
  .doc-preview em {{ font-style: italic; }}
</style>
<div class="doc-preview">{inner_html}</div>
"""


def _docx_to_basic_html(doc_path: str) -> str:
    """Fallback: extract text from DOCX and render as basic HTML."""
    try:
        doc = Document(doc_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                parts.append("<br/>")
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                parts.append(f"<h1>{_escape(text)}</h1>")
            elif "Heading 2" in style:
                parts.append(f"<h2>{_escape(text)}</h2>")
            elif "Heading 3" in style:
                parts.append(f"<h3>{_escape(text)}</h3>")
            else:
                # Check first run bold/italic
                bold = para.runs and para.runs[0].bold
                inner = f"<strong>{_escape(text)}</strong>" if bold else _escape(text)
                parts.append(f"<p>{inner}</p>")
        for table in doc.tables:
            rows_html = []
            for row in table.rows:
                cells = "".join(f"<td>{_escape(c.text)}</td>" for c in row.cells)
                rows_html.append(f"<tr>{cells}</tr>")
            parts.append(f"<table>{''.join(rows_html)}</table>")
        return _wrap_preview_html("".join(parts))
    except Exception as e:
        return f"<div>Preview error: {e}</div>"


def _escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def highlight_placeholders_html(template_docx_path: str) -> str:
    """
    Render the template DOCX as HTML with {{PLACEHOLDER}} tokens visually highlighted.
    """
    try:
        import mammoth
        with open(template_docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value
    except Exception:
        try:
            doc = Document(template_docx_path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(f"<p>{_escape(para.text)}</p>")
            for table in doc.tables:
                rows_html = []
                for row in table.rows:
                    cells = "".join(f"<td>{_escape(c.text)}</td>" for c in row.cells)
                    rows_html.append(f"<tr>{cells}</tr>")
                parts.append(f"<table>{''.join(rows_html)}</table>")
            html = "".join(parts)
        except Exception as e:
            return f"<div>Template preview error: {e}</div>"

    # Highlight {{PLACEHOLDER}} tokens
    def highlight(m):
        key = m.group(1)
        label = key.replace("_", " ").title()
        return (
            f'<mark style="background:#dbeafe;color:#1e40af;border:1px solid #93c5fd;'
            f'border-radius:4px;padding:1px 4px;font-weight:600;font-size:0.85em;">'
            f'{{{{{key}}}}}</mark>'
        )

    highlighted = re.sub(r'\{\{([A-Z0-9_]+)\}\}', highlight, html)
    return _wrap_preview_html(highlighted)


def get_placeholder_keys_from_docx(doc_path: str) -> list:
    """Extract all {{KEY}} placeholder keys from a DOCX file."""
    try:
        doc = Document(doc_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_parts.append(para.text)
        full_text = "\n".join(text_parts)
        return list(dict.fromkeys(re.findall(r'\{\{([A-Z0-9_]+)\}\}', full_text)))
    except Exception:
        return []


# ── PDF → preview image (using poppler pdftoppm) ────────────────────────────

def pdf_to_preview_images(pdf_path: str, output_dir: str, max_pages: int = 3) -> list:
    """
    Convert PDF pages to PNG images for browser preview.
    Returns list of output image paths.
    """
    os.makedirs(output_dir, exist_ok=True)
    base = os.path.join(output_dir, "page")
    try:
        import subprocess
        result = subprocess.run(
            ["pdftoppm", "-png", "-r", "120", "-l", str(max_pages), pdf_path, base],
            capture_output=True, timeout=30
        )
        images = sorted([
            os.path.join(output_dir, f)
            for f in os.listdir(output_dir)
            if f.startswith("page") and f.endswith(".png")
        ])
        return images
    except Exception as e:
        return []


def image_to_preview_html(image_path: str) -> str:
    """Return HTML to display an image as a document preview."""
    import base64
    try:
        with open(image_path, "rb") as f:
            data = base64.b64encode(f.read()).decode()
        ext = os.path.splitext(image_path)[1].lower().lstrip(".")
        mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png"}.get(ext, "image/png")
        return f'<div style="text-align:center;padding:8px;"><img src="data:{mime};base64,{data}" style="max-width:100%;border:1px solid #e2e8f0;border-radius:4px;" /></div>'
    except Exception as e:
        return f"<div>Image preview error: {e}</div>"
