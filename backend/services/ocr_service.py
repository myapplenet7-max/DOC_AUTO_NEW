"""
Improved OCR + variable detection pipeline.

Key improvements over v1:
- Detect document type first, then only suggest relevant fields
- ALL_CAPS proper noun detection (catches Indian legal name conventions)
- Multi-pass scanning — pass 2 ignores already-tagged values
- No fixed 25-field form: output is strictly what was detected
"""
import re
import os


# ── Text extraction (unchanged) ──────────────────────────────────────────────

def extract_text(file_path: str, preprocess_params: dict | None = None,
                 document_id: int | None = None, db=None) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        return _extract_from_txt(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    elif ext == ".odt":
        return _extract_from_odt(file_path)
    elif ext == ".pdf":
        text = _extract_from_pdf(file_path)
        if text and text.strip():
            return text
        return _extract_from_scanned_pdf(file_path, document_id=document_id, db=db)
    else:
        return _extract_from_image(file_path, preprocess_params=preprocess_params,
                                   document_id=document_id, db=db)


def _extract_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        return f"TXT extraction error: {str(e)}"


def _extract_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append("  |  ".join(row_texts))
        return "\n".join(parts)
    except Exception as e:
        return f"DOCX extraction error: {str(e)}"


def _extract_from_odt(file_path: str) -> str:
    try:
        from odf.opendocument import load
        from odf.text import P
        from odf import teletype
        doc = load(file_path)
        parts = []
        for para in doc.getElementsByType(P):
            text = teletype.extractText(para).strip()
            if text:
                parts.append(text)
        return "\n".join(parts)
    except ImportError:
        return _extract_odt_fallback(file_path)
    except Exception as e:
        return f"ODT extraction error: {str(e)}"


def _extract_odt_fallback(file_path: str) -> str:
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(file_path) as z:
            with z.open("content.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                texts = []
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        texts.append(elem.text.strip())
                return "\n".join(texts)
    except Exception as e:
        return f"ODT fallback error: {str(e)}"


def _ocr_pil_image(pil_image) -> str:
    """Run Tesseract on a PIL image — Telugu+English → English → bare fallback."""
    try:
        import pytesseract
    except ImportError:
        return "[OCR_UNAVAILABLE]"
    try:
        try:
            return pytesseract.image_to_string(pil_image, lang="tel+eng")
        except pytesseract.TesseractNotFoundError:
            return "[OCR_UNAVAILABLE]"
        except Exception:
            try:
                return pytesseract.image_to_string(pil_image, lang="eng")
            except pytesseract.TesseractNotFoundError:
                return "[OCR_UNAVAILABLE]"
            except Exception:
                return pytesseract.image_to_string(pil_image)
    except Exception as e:
        import logging
        logging.getLogger(__name__).error("pytesseract error: %s", e)
        return "[OCR_UNAVAILABLE]"


def _extract_from_image(file_path: str, preprocess_params: dict | None = None,
                        document_id: int | None = None, db=None) -> str:
    import logging
    logger = logging.getLogger(__name__)

    try:
        from services.image_preprocess import preprocess_image, image_hash
        processed = preprocess_image(file_path, params=preprocess_params)
        logger.info("Preprocessing applied to %s", file_path)
    except Exception as e:
        logger.warning("Preprocessing failed (%s) — falling back to raw image", e)
        try:
            from PIL import Image
            processed = Image.open(file_path)
        except Exception as e2:
            logger.error("Failed to open image %s: %s", file_path, e2)
            return f"[Image open error: {e2}]"

    text = _ocr_pil_image(processed)

    # Silently save training pair for future fine-tuning (data collection only)
    if db is not None and text and text != "[OCR_UNAVAILABLE]":
        try:
            import json
            from services.image_preprocess import image_hash as _hash
            from models import OCRTrainingData
            img_hash = _hash(file_path)
            existing = db.query(OCRTrainingData).filter(
                OCRTrainingData.image_hash == img_hash
            ).first()
            if not existing:
                pair = OCRTrainingData(
                    document_id=document_id,
                    image_hash=img_hash,
                    preprocessing_params=json.dumps(preprocess_params or {}),
                    ocr_output_text=text[:4000],
                )
                db.add(pair)
                db.commit()
        except Exception as e:
            logger.debug("Training pair save skipped: %s", e)

    return text


def _extract_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        return f"PDF extraction error: {str(e)}"


def _extract_from_scanned_pdf(file_path: str, document_id: int | None = None,
                              db=None) -> str:
    try:
        from pdf2image import convert_from_path
        pages = convert_from_path(file_path, dpi=300)
        text_parts = []
        for page_image in pages:
            text_parts.append(_ocr_pil_image(page_image))
        return "\n".join(text_parts)
    except Exception as e:
        return f"Scanned PDF OCR error: {str(e)}"


def _fallback_image_text_from_pil(pil_image):
    return _ocr_pil_image(pil_image)


# ── Document type detection ──────────────────────────────────────────────────

DOC_TYPE_SPECS = {
    "birth_affidavit": {
        "display": "Birth Affidavit",
        "emoji": "👶",
        "keywords": ["affidavit for date of birth", "birth affidavit", "date of birth", "born", "delivery", "birth registration", "child", "municipality"],
        "field_keys": ["deponent_name", "father_name", "mother_name", "child_name", "date_of_birth", "place_of_birth", "age", "municipality", "address", "district", "state", "mandal", "door_number"],
    },
    "sale_deed": {
        "display": "Sale Deed",
        "emoji": "🏠",
        "keywords": ["sale deed", "vikaraya", "sale consideration", "vendor", "vendee", "విక్రయ దస్తావేజు", "seller", "buyer", "purchase price"],
        "field_keys": ["seller_name", "buyer_name", "witness_name", "survey_number", "door_number", "boundaries", "village", "mandal", "district", "state", "sale_amount", "registration_date", "registration_number", "property_area", "address"],
    },
    "rental_agreement": {
        "display": "Rental Agreement",
        "emoji": "🔑",
        "keywords": ["rental agreement", "lease agreement", "tenant", "landlord", "monthly rent", "license fee", "lessor", "lessee"],
        "field_keys": ["landlord_name", "tenant_name", "property_address", "monthly_rent", "start_date", "end_date", "security_deposit", "village", "district", "state"],
    },
    "gpa": {
        "display": "General Power of Attorney",
        "emoji": "📋",
        "keywords": ["general power of attorney", "power of attorney", " gpa ", "attorney in fact"],
        "field_keys": ["grantor_name", "grantee_name", "grantor_address", "property_address", "survey_number", "village", "district", "state", "date"],
    },
    "affidavit": {
        "display": "Affidavit",
        "emoji": "📜",
        "keywords": ["affidavit", "solemnly affirm", "deponent", "sworn"],
        "field_keys": ["deponent_name", "father_name", "address", "village", "mandal", "district", "state", "date", "age"],
    },
    "court_petition": {
        "display": "Court Petition",
        "emoji": "⚖️",
        "keywords": ["petition", "honorable court", "high court", "respondent", "petitioner", "the registrar", "case no"],
        "field_keys": ["petitioner_name", "respondent_name", "case_number", "court_name", "advocate_name", "date", "address"],
    },
    "gift_deed": {
        "display": "Gift Deed",
        "emoji": "🎁",
        "keywords": ["gift deed", "gifted", "donor", "donee", "gift of property"],
        "field_keys": ["donor_name", "donee_name", "relationship", "property_address", "survey_number", "village", "district", "state", "date"],
    },
    "legal_notice": {
        "display": "Legal Notice",
        "emoji": "📢",
        "keywords": ["legal notice", "notice", "demand notice", "advocate notice"],
        "field_keys": ["sender_name", "recipient_name", "advocate_name", "subject", "demand_amount", "address", "district", "state", "date"],
    },
    "certificate": {
        "display": "Certificate",
        "emoji": "🎓",
        "keywords": ["certificate", "certified", "this is to certify", "issued to"],
        "field_keys": ["recipient_name", "father_name", "purpose", "issuing_authority", "date", "village", "district", "state"],
    },
    "partition_deed": {
        "display": "Partition Deed",
        "emoji": "📄",
        "keywords": ["partition deed", "partition", "share", "co-owner", "panchayat partition"],
        "field_keys": ["party_1_name", "party_2_name", "survey_number", "boundaries", "village", "district", "state", "date"],
    },
}


def detect_document_type(raw_text: str) -> dict:
    """Classify the document and return type metadata."""
    text_lower = raw_text.lower()

    best_type = None
    best_score = 0

    for type_id, spec in DOC_TYPE_SPECS.items():
        score = 0
        for kw in spec["keywords"]:
            if kw in text_lower:
                score += (2 if len(kw) > 10 else 1)
        if score > best_score:
            best_score = score
            best_type = type_id

    if best_type:
        spec = DOC_TYPE_SPECS[best_type]
        return {
            "type": best_type,
            "display": spec["display"],
            "emoji": spec["emoji"],
            "confidence": min(best_score / 3, 1.0),
            "field_keys": spec["field_keys"],
        }

    return {
        "type": "other",
        "display": "Document",
        "emoji": "📄",
        "confidence": 0.3,
        "field_keys": ["name", "address", "date", "district", "state"],
    }


# ── Variable Detection ────────────────────────────────────────────────────────

# Words to exclude from ALL_CAPS name detection
_CAPS_EXCLUDE = {
    "I", "A", "AN", "THE", "OF", "IN", "AT", "BY", "TO", "AS", "OR", "AND",
    "BE", "IS", "ARE", "WAS", "NOT", "DO", "NO", "MY", "HE", "SHE", "WE",
    "IT", "ON", "UP", "IF", "SO", "GO", "HIS", "HER", "ITS", "FOR", "FROM",
    "WITH", "THIS", "THAT", "WILL", "BEEN", "HAVE", "HAD", "HAS", "CAN",
    "AGED", "ABOUT", "YEARS", "AGE", "YEAR", "DAY", "DATE", "DAYS",
    "AFFIDAVIT", "DEED", "AGREEMENT", "NOTICE", "CERTIFICATE", "PETITION",
    "SALE", "BIRTH", "RENTAL", "GIFT", "PARTITION", "COURT", "HIGH",
    "DELIVERY", "REGISTRY", "REVENUE", "GOVERNMENT", "DISTRICT", "MANDAL",
    "VILLAGE", "STATE", "INDIA", "ANDHRA", "TELANGANA", "PRADESH",
    "SIR", "MADAM", "THE", "HEREBY", "SOLEMNLY", "AFFIRM", "DECLARE",
    "CONTENTS", "ABOVE", "TRUE", "CORRECT", "NOTHING", "CONCEALED",
    "SUBMITTED", "RESPECTFULLY", "HUMBLY", "REQUESTED", "KINDLY",
    "REGISTRATION", "AUTHORITY", "AUTHORITY", "MUNICIPAL", "MUNICIPALITY",
    "DIVISIONAL", "OFFICER", "DIVISION", "SUB", "REGISTRAR", "TALUK",
    "PLEASE", "THANK", "YOURS", "FAITHFULLY", "SINCERELY", "WITNESS",
    "SCHEDULE", "ANNEXURE", "EXHIBIT", "PARA", "CLAUSE", "SECTION",
    "NO", "NUMBER", "NOS", "SL", "SR", "TOTAL", "AMOUNT", "RUPEES",
    "RS", "PAGE", "OF", "PART", "JOINT", "DEED", "DOCUMENT",
    "REGISTER", "STAMP", "DUTY", "FEE", "PENALTY", "FINE",
    "RD", "TH", "ST", "ND",  # ordinal suffixes
    "FLOOR", "GROUND", "FLAT", "APARTMENT", "BLOCK",
    "H", "NO", "D", "SY", "SUR", "W", "S",  # abbreviations alone
    "MALE", "FEMALE", "MINOR", "MAJOR",
    "COOLIE", "WORKER", "FARMER", "LABOURER", "TEACHER", "OFFICER",
    "WORKING", "RESIDING", "RESIDENT", "LIVING", "STAYING",
    "HEREBY", "THEREFORE", "WHEREFORE", "WHEREAS", "HEREIN",
    "ABOVE", "BELOW", "SAID", "SAME", "SUCH", "EACH", "BOTH",
    "THREE", "TWO", "ONE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE",
    "TEN", "ELEVEN", "TWELVE", "LAKH", "THOUSAND", "HUNDRED", "CRORE",
    "FIRST", "SECOND", "THIRD", "FOURTH", "LAST", "NEXT", "ONLY",
}


def _detect_all_caps_phrases(raw_text: str, already_found: set = None) -> list:
    """
    Find sequences of ALL_CAPS words likely to be proper names/locations.
    Returns list of {text, suggested_key, category} dicts.
    """
    if already_found is None:
        already_found = set()

    # Match sequences of 2+ uppercase words (at least 2 consecutive)
    pattern = r'\b([A-Z]{2,}(?:\s+[A-Z]{2,})*)\b'
    candidates = []
    seen = set()

    for match in re.finditer(pattern, raw_text):
        phrase = match.group(1)
        words = phrase.split()

        # Filter individual excluded words
        filtered = [w for w in words if w not in _CAPS_EXCLUDE and len(w) > 1]
        if not filtered:
            continue

        # Rebuild phrase from filtered words (take longest consecutive run)
        phrase_filtered = " ".join(filtered)
        if len(phrase_filtered) < 3:
            continue

        # Skip if already converted to placeholder
        if phrase_filtered in already_found or phrase in already_found:
            continue

        # Skip very common standalone words
        if phrase_filtered in _CAPS_EXCLUDE:
            continue

        # Deduplicate
        if phrase_filtered in seen:
            continue
        seen.add(phrase_filtered)

        # Guess category and suggested key
        suggested_key, category = _guess_name_key(phrase_filtered)
        candidates.append({
            "text": phrase_filtered,
            "suggested_key": suggested_key,
            "category": category,
        })

    return candidates


def _guess_name_key(phrase: str) -> tuple:
    """Guess the variable key and category for a detected phrase."""
    p = phrase.upper()
    words = phrase.split()

    # Location patterns in the phrase itself
    if any(x in p for x in ["NAGAR", "PURAM", "PALLI", "PETA", "GUDA", "PADU", "PALLE", "COLONY", "STREET"]):
        return "LOCATION", "Location"
    if any(x in p for x in ["DISTRICT", "MANDAL", "TALUK", "MUNICIPALITY", "CORPORATION"]):
        return "DISTRICT", "Location"
    if any(x in p for x in ["ANDHRA", "TELANGANA", "PRADESH", "KARNATAKA", "MAHARASHTRA", "ODISHA", "KERALA"]):
        return "STATE", "Location"

    # Multi-word → likely person name
    if len(words) >= 2:
        return "PERSON_NAME", "Person Details"

    # Single word: could be village/town or surname — stay ambiguous, context will resolve
    return "PROPER_NAME", "General"


def detect_fields(raw_text: str) -> dict:
    """
    Returns a dict with raw_text + doc_type info. Actual field detection
    is now done in detect_placeholders() via comprehensive scanning.
    """
    doc_type_info = detect_document_type(raw_text)
    return {
        "raw_text": raw_text,
        "doc_type": doc_type_info["type"],
        "doc_type_display": doc_type_info["display"],
        "doc_type_emoji": doc_type_info["emoji"],
        "doc_type_confidence": doc_type_info["confidence"],
    }


def detect_placeholders(raw_text: str, detected_fields: dict = None) -> list:
    """
    Comprehensive placeholder detection:
    1. Regex patterns for IDs, numbers, dates, amounts
    2. ALL_CAPS proper noun detection
    3. Keyword-anchor based detection for common field patterns
    Returns list of placeholder dicts.
    """
    placeholders = []
    seen_values = set()

    def add(key, placeholder, value, confidence, category):
        value = str(value).strip()
        if not value or len(value) < 2:
            return
        # Avoid duplicates by value or placeholder name
        for p in placeholders:
            if p["value"] == value:
                return
            # Allow same placeholder key only if value differs AND placeholder name is unique
            if p["placeholder"] == placeholder:
                # Same placeholder name — skip (first detection wins)
                return
        if value in seen_values:
            return
        seen_values.add(value)
        placeholders.append({
            "key": key,
            "placeholder": placeholder,
            "value": value,
            "confidence": confidence,
            "category": category,
            "approved": confidence >= 0.85,
        })

    lines = raw_text.split("\n")

    # ── Pass 1: Regex patterns ────────────────────────────────────────────────

    for line in lines:
        # Aadhaar: 12 digits
        m = re.search(r'\b(\d{4}\s?\d{4}\s?\d{4})\b', line)
        if m:
            add("aadhar_number", "AADHAAR_NUMBER", m.group(1).replace(" ", ""), 0.99, "Identity")

        # PAN
        m = re.search(r'\b([A-Z]{5}[0-9]{4}[A-Z])\b', line)
        if m:
            add("pan_number", "PAN_NUMBER", m.group(1), 0.99, "Identity")

        # Mobile
        m = re.search(r'\b([6-9]\d{9})\b', line)
        if m:
            add("mobile", "MOBILE_NUMBER", m.group(1), 0.95, "Contact")

        # Email
        m = re.search(r'[\w.\-]+@[\w.\-]+\.\w+', line)
        if m:
            add("email", "EMAIL_ADDRESS", m.group(0), 0.95, "Contact")

        # PIN code
        m = re.search(r'\b([1-9]\d{5})\b', line)
        if m and not re.search(r'\b\d{4}\s?\d{4}\s?\d{4}\b', line):
            add("pincode", "PINCODE", m.group(1), 0.90, "Location")

        # Dates — various formats
        for date_pat in [
            r'\b(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{4})\b',
            r'\b(\d{1,2}(?:st|nd|rd|th)\s+(?:day\s+of\s+)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})\b',
            r'\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b',
            r'\b((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4})\b',
        ]:
            m = re.search(date_pat, line, re.IGNORECASE)
            if m:
                val = m.group(1)
                # Classify as DOB or registration date based on context
                context = line.lower()
                if any(w in context for w in ["born", "birth", "dob", "date of birth"]):
                    add("date_of_birth", "DATE_OF_BIRTH", val, 0.92, "Dates")
                elif any(w in context for w in ["register", "registration", "executed", "registered"]):
                    add("registration_date", "REGISTRATION_DATE", val, 0.92, "Dates")
                else:
                    add("date", "DATE", val, 0.85, "Dates")
                break

        # Amount / money
        m = re.search(r'(?:rs\.?|₹|రూ\.?)\s*([\d,]+(?:\.\d+)?)', line, re.IGNORECASE)
        if m:
            context = line.lower()
            if any(w in context for w in ["sale", "purchase", "consideration"]):
                add("sale_amount", "SALE_AMOUNT", m.group(1), 0.90, "Financial")
            elif any(w in context for w in ["rent", "monthly", "lease"]):
                add("monthly_rent", "MONTHLY_RENT", m.group(1), 0.90, "Financial")
            else:
                add("amount", "AMOUNT", m.group(1), 0.80, "Financial")

        # House / door number
        m = re.search(r'\b(?:H\.?No\.?|D\.?No\.?|Door\s*No\.?|House\s*No\.?)\s*([\d\/\-A-Za-z]+)', line, re.IGNORECASE)
        if m:
            add("door_number", "DOOR_NUMBER", m.group(1), 0.90, "Location")

        # Survey number
        m = re.search(r'\b(?:Sy\.?\s*No\.?|Survey\s*No\.?|S\.?\s*No\.?)\s*([\d\/\-A-Za-z]+)', line, re.IGNORECASE)
        if m:
            add("survey_number", "SURVEY_NUMBER", m.group(1), 0.90, "Property")

        # Registration / document number
        m = re.search(r'\b(?:Reg\.?\s*No\.?|Document\s*No\.?|Doc\.?\s*No\.?)\s*([\d\/\-A-Za-z]+)', line, re.IGNORECASE)
        if m:
            add("registration_number", "REGISTRATION_NUMBER", m.group(1), 0.88, "Legal")

        # Case number
        m = re.search(r'\b(?:Case\s*No\.?|O\.?P\.?\s*No\.?|C\.?C\.?\s*No\.?|W\.?P\.?\s*No\.?|Petition\s*No\.?)\s*([\d\/\-A-Za-z]+)', line, re.IGNORECASE)
        if m:
            add("case_number", "CASE_NUMBER", m.group(1), 0.90, "Legal")

        # Age
        m = re.search(r'\b(?:aged?|age)\s+(?:about\s+)?(\d{1,3})\s*(?:years?|yrs?)?\b', line, re.IGNORECASE)
        if m:
            add("age", "AGE", m.group(1), 0.85, "Person Details")

    # ── Pass 2: Keyword-anchor extraction ────────────────────────────────────

    ANCHORS = [
        # (field_key, placeholder, category, confidence, keywords)
        ("deponent_name",   "DEPONENT_NAME",   "Person Details", 0.88, ["deponent", "i, the deponent", "i am the deponent"]),
        ("father_name",     "FATHER_NAME",      "Person Details", 0.88, ["s/o", "son of", "father:", "father's name", "తండ్రి"]),
        ("mother_name",     "MOTHER_NAME",      "Person Details", 0.88, ["w/o", "wife of", "mother:", "my wife namely", "m/o"]),
        ("husband_name",    "HUSBAND_NAME",     "Person Details", 0.85, ["husband", "w/o", "భర్త"]),
        ("child_name",      "CHILD_NAME",       "Person Details", 0.88, ["named", "child named", "son named", "daughter named", "child's name", "named by"]),
        ("advocate_name",   "ADVOCATE_NAME",    "Legal",          0.85, ["through advocate", "through counsel", "advocate:", "counsel:"]),
        ("place_of_birth",  "PLACE_OF_BIRTH",   "Location",       0.85, ["born at", "place of birth", "birth place", "delivery at"]),
        ("municipality",    "MUNICIPALITY",     "Location",       0.85, ["municipality of", "municipality:", "municipal corporation", "panchayat", "gram panchayat"]),
        ("village",         "VILLAGE",          "Location",       0.85, ["village:", "vill.", "village of", "గ్రామం", "gram:"]),
        ("mandal",          "MANDAL",           "Location",       0.85, ["mandal:", "మండలం", "mandal of", "taluk"]),
        ("district",        "DISTRICT",         "Location",       0.85, ["district:", "dist:", "జిల్లా"]),
        ("state",           "STATE",            "Location",       0.80, ["state:", "రాష్ట్రం", "state of"]),
        ("boundaries",      "BOUNDARIES",       "Property",       0.80, ["bounded by", "east:", "north:", "south:", "west:", "నాలుగు హద్దులు"]),
        ("designation",     "DESIGNATION",      "Person Details", 0.80, ["designation:", "working as", "occupation:", "profession:"]),
    ]

    for line in lines:
        line_lower = line.lower()
        for field_key, placeholder, category, confidence, keywords in ANCHORS:
            if any(k in line_lower or k in line for k in keywords):
                # Extract the value after the keyword/separator
                val = _extract_value_after_keyword(line, keywords)
                if val and len(val) > 1:
                    add(field_key, placeholder, val, confidence, category)

    # ── Pass 3: ALL_CAPS proper noun detection ────────────────────────────────

    already_found_values = {p["value"] for p in placeholders}
    caps_candidates = _detect_all_caps_phrases(raw_text, already_found_values)

    for candidate in caps_candidates:
        text = candidate["text"]
        category = candidate["category"]
        # Assign a more specific key based on context
        suggested_key = _contextual_name_key(text, raw_text, placeholders, category)
        # Confidence based on phrase length (longer = more likely real name)
        conf = 0.88 if len(text.split()) >= 2 else 0.78

        add(suggested_key.lower(), suggested_key, text, conf, category)

    # ── Final dedup pass: case-insensitive name, keep longer value ────────────
    # Safety net: collapse any remaining duplicates that slipped through the
    # per-call checks above (can happen when the same placeholder name is
    # emitted by two different detection passes with slightly different values).
    seen_ph: dict[str, int] = {}  # normalised_placeholder_name → index in final list
    deduped: list = []
    for ph in placeholders:
        norm = ph["placeholder"].upper()
        if norm in seen_ph:
            existing_idx = seen_ph[norm]
            existing = deduped[existing_idx]
            # Replace with the longer / higher-confidence value
            if (len(ph["value"]) > len(existing["value"]) or
                    ph["confidence"] > existing["confidence"]):
                deduped[existing_idx] = ph
        else:
            seen_ph[norm] = len(deduped)
            deduped.append(ph)

    if len(deduped) < len(placeholders):
        import logging as _log
        _log.getLogger("docauto.ocr").info(
            "detect_placeholders: final dedup removed %d duplicate(s) (%d → %d)",
            len(placeholders) - len(deduped), len(placeholders), len(deduped),
        )

    return deduped


def _contextual_name_key(phrase: str, full_text: str, existing: list, category: str) -> str:
    """Assign a contextual variable key based on surrounding text."""
    # Check context around this phrase in the full text
    idx = full_text.find(phrase)
    context = ""
    if idx >= 0:
        context = full_text[max(0, idx - 60):idx + len(phrase) + 30].lower()

    # Check what we've already assigned
    existing_keys = {p["placeholder"] for p in existing}

    if any(w in context for w in ["named", "child named", "son named", "daughter named", "named by"]):
        return "CHILD_NAME" if "CHILD_NAME" not in existing_keys else "CHILD_NAME_2"
    if any(w in context for w in ["s/o", "son of", "w/o", "wife of", "father", "mother"]):
        if "FATHER_NAME" not in existing_keys:
            return "FATHER_NAME"
        if "MOTHER_NAME" not in existing_keys:
            return "MOTHER_NAME"
    if any(w in context for w in ["deponent", "i,", "i am"]):
        return "DEPONENT_NAME" if "DEPONENT_NAME" not in existing_keys else "PERSON_NAME"
    if any(w in context for w in ["seller", "vendor", "grantor"]):
        return "SELLER_NAME" if "SELLER_NAME" not in existing_keys else "PERSON_NAME"
    if any(w in context for w in ["buyer", "vendee", "grantee"]):
        return "BUYER_NAME" if "BUYER_NAME" not in existing_keys else "PERSON_NAME"
    if any(w in context for w in ["landlord", "lessor", "owner"]):
        return "LANDLORD_NAME" if "LANDLORD_NAME" not in existing_keys else "PERSON_NAME"
    if any(w in context for w in ["tenant", "lessee", "occupant"]):
        return "TENANT_NAME" if "TENANT_NAME" not in existing_keys else "PERSON_NAME"

    # Location words in phrase itself
    p_upper = phrase.upper()
    if any(x in p_upper for x in ["NAGAR", "PURAM", "PALLI", "PETA"]):
        return "LOCATION_NAME"
    if category == "Location":
        if "VILLAGE" not in existing_keys:
            return "VILLAGE"
        if "DISTRICT" not in existing_keys:
            return "DISTRICT"
        return "LOCATION_NAME"

    # Generic person name if not resolved
    if category == "Person Details":
        if "PERSON_NAME" not in existing_keys:
            return "PERSON_NAME"
        # Increment
        for i in range(2, 10):
            k = f"PERSON_NAME_{i}"
            if k not in existing_keys:
                return k

    return phrase.replace(" ", "_")[:30]


def _extract_value_after_keyword(line: str, keywords: list) -> str:
    """Extract the value following a keyword in a line (stops at sentence breaks)."""
    line_lower = line.lower()
    for kw in keywords:
        found_idx = -1
        if kw in line_lower:
            found_idx = line_lower.index(kw)
        elif kw in line:
            found_idx = line.index(kw)

        if found_idx >= 0:
            rest = line[found_idx + len(kw):].strip()
            rest = re.sub(r'^[:\-–\s]+', '', rest).strip()
            if not rest:
                continue
            # Truncate at known stop patterns
            stop_patterns = [
                r'\s+[Aa]ged?\b',        # "Aged about"
                r'\s+[Ss]/[oO]\b',       # "S/o"
                r'\s+[Ww]/[oO]\b',       # "W/o"
                r'\s+on\s+\d',           # "on 18-09-2008"
                r'\s+in\s+\d',           # "in 3RD"
                r'\s+at\s+[A-Z]',        # "at H.No"
                r'\s+[Rr]esiding\b',
                r'\s+[Dd]elivered\b',
                r',\s*[Aa]ndhra\b',
                r',\s*[Tt]elangana\b',
                r'\s*\.(?:\s|$)',         # sentence period
            ]
            truncated = rest
            for pat in stop_patterns:
                m = re.search(pat, truncated)
                if m:
                    truncated = truncated[:m.start()].strip()

            # Strip leading function words (by, the, of, a, an)
            truncated = re.sub(r'^(?:by|the|of|a|an|to|from|for)\s+', '', truncated, flags=re.IGNORECASE).strip()

            # Also limit to max 6 words
            words = truncated.split()
            if len(words) > 6:
                truncated = " ".join(words[:6])

            truncated = truncated.rstrip(",;").strip()
            if truncated and len(truncated) > 1:
                return truncated[:150]
    return ""


# ── Second pass / rescan ──────────────────────────────────────────────────────

def rescan_for_missing(raw_text: str, existing_placeholders: list) -> list:
    """
    Second pass: re-run detection ignoring already-tagged values.
    Returns NEW placeholders not already in existing_placeholders.

    Uses case-insensitive name comparison so that e.g. "AMOUNT" from the
    first pass correctly blocks "amount" (or vice-versa) from being returned
    as a duplicate by the rescan.
    """
    existing_values = {p["value"] for p in existing_placeholders}
    # Normalise to uppercase for case-insensitive name comparison
    existing_ph_names_upper = {p["placeholder"].upper() for p in existing_placeholders}

    # Mask already-detected values in the text so they don't trigger again
    text_for_scan = raw_text
    for val in sorted(existing_values, key=len, reverse=True):
        if val and len(val) > 2:
            text_for_scan = text_for_scan.replace(val, "____ALREADY_TAGGED____")

    # Run detection on the masked text
    new_ph = detect_placeholders(text_for_scan, {})

    # Filter: skip if name already exists (case-insensitive) or value is masked
    results = []
    for ph in new_ph:
        if ph["placeholder"].upper() in existing_ph_names_upper:
            continue
        if "ALREADY_TAGGED" in ph["value"]:
            continue
        results.append(ph)

    return results


# ── Template generation ───────────────────────────────────────────────────────

def generate_template_from_text(raw_text: str, approved_placeholders: list) -> str:
    """Replace detected values with {{PLACEHOLDER}} tokens in raw text."""
    template = raw_text
    sorted_ph = sorted(approved_placeholders, key=lambda x: len(x.get("value", "")), reverse=True)
    for ph in sorted_ph:
        value = ph.get("value", "").strip()
        key = ph.get("key", "")
        placeholder_name = ph.get("placeholder", key.upper() if key else "VALUE")
        if value and placeholder_name:
            template = template.replace(value, f"{{{{{placeholder_name}}}}}")
    return template
