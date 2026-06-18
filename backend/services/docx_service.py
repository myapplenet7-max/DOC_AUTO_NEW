from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime


def _set_telugu_capable_font(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Nirmala UI"
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Nirmala UI")
    rfonts.set(qn("w:hAnsi"), "Nirmala UI")
    rfonts.set(qn("w:cs"), "Nirmala UI")


def generate_docx(fields: dict, output_path: str):
    doc = Document()
    _set_telugu_capable_font(doc)

    title = doc.add_heading("DOCUMENT RECORD", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_paragraph("")

    # Person Details
    person_data = {
        "Full Name":       fields.get("full_name", ""),
        "Father's Name":   fields.get("father_name", ""),
        "Husband's Name":  fields.get("husband_name", ""),
        "Deponent Name":   fields.get("deponent_name", ""),
        "Advocate":        fields.get("advocate_name", ""),
        "Date of Birth":   fields.get("date_of_birth", ""),
        "Aadhaar Number":  fields.get("aadhar_number", ""),
        "PAN Number":      fields.get("pan_number", ""),
        "Mobile":          fields.get("mobile", ""),
        "Email":           fields.get("email", ""),
        "Pincode":         fields.get("pincode", ""),
    }
    filled_person = {k: v for k, v in person_data.items() if v}
    if filled_person:
        doc.add_heading("Person / Applicant Details", level=2)
        _add_field_table(doc, filled_person)
        doc.add_paragraph("")

    # Location
    location_data = {
        "Address":   fields.get("address", ""),
        "Village":   fields.get("village", ""),
        "Mandal":    fields.get("mandal", ""),
        "District":  fields.get("district", ""),
        "State":     fields.get("state", ""),
        "Pincode":   fields.get("pincode", ""),
    }
    filled_location = {k: v for k, v in location_data.items() if v and k not in filled_person}
    if filled_location:
        doc.add_heading("Location Details", level=2)
        _add_field_table(doc, filled_location)
        doc.add_paragraph("")

    # Property Details
    property_data = {
        "Survey Number":       fields.get("survey_number", ""),
        "Door Number":         fields.get("door_number", ""),
        "Plot Number":         fields.get("plot_number", ""),
        "Property Details":    fields.get("property_details", ""),
        "Boundaries":          fields.get("boundaries", ""),
        "Sale Amount (₹)":     fields.get("sale_amount", ""),
        "Registration Date":   fields.get("registration_date", ""),
        "Registration Number": fields.get("registration_number", ""),
    }
    filled_property = {k: v for k, v in property_data.items() if v}
    if filled_property:
        doc.add_heading("Property Details", level=2)
        _add_field_table(doc, filled_property)
        doc.add_paragraph("")

    # Legal Details
    legal_data = {
        "Court Case Number": fields.get("court_case_number", ""),
        "Advocate":          fields.get("advocate_name", ""),
    }
    filled_legal = {k: v for k, v in legal_data.items() if v and k not in filled_person}
    if filled_legal:
        doc.add_heading("Legal Details", level=2)
        _add_field_table(doc, filled_legal)
        doc.add_paragraph("")

    # Raw text (small)
    raw = fields.get("raw_text", "")
    if raw:
        doc.add_paragraph("")
        doc.add_heading("Raw Extracted Text", level=3)
        p = doc.add_paragraph(raw)
        if p.runs:
            p.runs[0].font.size = Pt(8)

    doc.save(output_path)


def _add_field_table(doc: Document, data: dict):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in data.items():
        if not value:
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value)
        for para in row[0].paragraphs:
            for run in para.runs:
                run.bold = True
